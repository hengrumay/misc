"""Seed synthetic study-protocol PDFs/DOCX into the ads_raw.protocols volume.

These are stand-in "redacted protocols" so the PDF/DOCX -> ai_extract path can be
demonstrated end-to-end before real redacted protocols are available +
SME-labeled expected fields. When real files arrive, just drop them into the
volume — no code change needed.

The codes embedded here are drawn from the *actual* synthetic RWD code universe
(ads_serving.patient_timeline) so each extracted cohort is non-empty:
  dx present:  E11.9 (T2DM), I10 (HTN), I50.9 (HF), I25.10 (CAD), N18.3 (CKD), K21.9, E03.9
  rx present:  00093-5117-16, 00054-0165-24 (top NDCs, ~800 patients each)

Runs on serverless (writes locally then copies to the UC volume, which is not
seekable). Resolves the volume path via lib/config.py.

CLI: python scripts/seed_protocol_pdfs.py
"""
from __future__ import annotations

# Serverless spark_python_task runs this file via exec() with no __file__ in
# globals; recover it from the frame so downstream Path(__file__) works.
try:
    __file__
except NameError:  # pragma: no cover
    import inspect as _inspect
    __file__ = _inspect.getfile(_inspect.currentframe())

import os
import shutil
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
from lib.config import cfg

# ---- protocol fixtures (realistic codes present in the synthetic RWD) --------
PROTOCOLS = {
    "poc_low": """# Study Protocol: Simple Prevalence Cohort (poc_low)

## Study Identification
- Study ID: poc_low
- Title: Prevalence of Type 2 Diabetes in a Managed Care Population
- Complexity: Low

## Objective
To estimate the prevalence of Type 2 Diabetes Mellitus among adult patients in a managed care population during the study period.

## Population
- Inclusion Criteria:
  - Continuous enrollment for the entire study period
  - Age 18-85 years at study index date
- Exclusion Criteria:
  - Other/unspecified hypothyroidism (ICD-10-CM: E03.9)

## Index Event
- Index Date: First occurrence of Type 2 Diabetes diagnosis (ICD-10-CM: E11.9) on or after January 1, 2018

## Covariates
- Hypertension (ICD-10-CM: I10)

## Study Period
- Start Date: 2018-01-01
- End Date: 2024-12-31

## Study Outcomes
- Primary Outcome: First heart failure event (ICD-10-CM: I50.9), follow-up 12 months
""",
    "poc_med": """# Study Protocol: Drug-Exposure New-User Cohort (poc_med)

## Study Identification
- Study ID: poc_med
- Title: Real-World Effectiveness and Safety of ACE Inhibitors in Hypertension Management
- Complexity: Medium

## Objective
To evaluate the real-world effectiveness and safety of ACE inhibitors versus other antihypertensive agents in patients with newly diagnosed hypertension.

## Population
- Inclusion Criteria:
  - Age 18-75 years at cohort entry
  - Continuous enrollment for 12 months before and 12 months after index event
  - No prior antihypertensive medication use in the 12-month baseline period (new-user cohort)
- Exclusion Criteria:
  - Chronic kidney disease stage 3 (ICD-10-CM: N18.3)

## Index Event
- Index Date: First pharmacy fill of the study antihypertensive agent (NDC 00093-5117-16)
- Washout: 365 days with no prior fill of the study drug

## Covariates (Baseline = 365 days pre-index)
- Diabetes (ICD-10-CM: E11.9)
- Heart failure (ICD-10-CM: I50.9)

## Study Period
- Start Date: 2018-01-01
- End Date: 2024-12-31
- Follow-up Duration: 12 months post-index

## Study Outcomes
- Primary Outcome: First heart failure hospitalization (ICD-10-CM: I50.9), follow-up 365 days
""",
    "poc_high": """# Study Protocol: Comparative Outcomes with Time-Varying Exposure (poc_high)

## Study Identification
- Study ID: poc_high
- Title: Comparative Effectiveness and Safety of Statin Intensification in Secondary Prevention of Cardiovascular Disease
- Complexity: High

## Objective
To compare high-intensity vs. moderate-intensity statin therapy in patients with established cardiovascular disease, evaluating time-varying exposure.

## Population
- Inclusion Criteria:
  - Age 40-75 years at cohort entry
  - Currently on a statin medication at baseline (NDC 00054-0165-24)
  - Continuous enrollment for 12 months before and 24 months after cohort entry
- Exclusion Criteria:
  - Chronic kidney disease stage 3 (ICD-10-CM: N18.3)
  - Gastro-esophageal reflux disease (ICD-10-CM: K21.9)

## Index Event and Exposure Definition
- Index Date: First statin fill (NDC 00054-0165-24)
- Time-Varying Exposure: exposure eras collapsing consecutive fills with up to 30-day gaps
- Washout: 365 days

## Covariates (Baseline = 365 days pre-index)
- Diabetes (ICD-10-CM: E11.9)
- Hypertension (ICD-10-CM: I10)
- Heart failure (ICD-10-CM: I50.9)

## Study Period
- Start Date: 2018-01-01
- End Date: 2024-12-31
- Follow-up Duration: 24 months post-index event

## Study Outcomes
- Primary Outcome: Heart failure event (ICD-10-CM: I50.9), follow-up 730 days
""",
}


def _flowables(md, styles):
    from reportlab.platypus import Paragraph, Spacer
    out = []
    for raw in md.splitlines():
        line = raw.rstrip()
        if not line.strip():
            out.append(Spacer(1, 4)); continue
        if line.startswith("## "):
            out.append(Paragraph(line[3:], styles["h2"]))
        elif line.startswith("# "):
            out.append(Paragraph(line[2:], styles["h1"]))
        elif line.strip().startswith("- "):
            ind = len(line) - len(line.lstrip())
            out.append(Paragraph("&bull; " + line.strip()[2:], styles["b2"] if ind >= 2 else styles["b1"]))
        else:
            out.append(Paragraph(line, styles["body"]))
    return out


def _write_pdf(local_path, md):
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate
    from reportlab.lib.units import inch
    base = getSampleStyleSheet()
    styles = {
        "h1": ParagraphStyle("h1", parent=base["Heading1"], fontSize=15, spaceAfter=8),
        "h2": ParagraphStyle("h2", parent=base["Heading2"], fontSize=12, spaceBefore=8, spaceAfter=4),
        "body": ParagraphStyle("body", parent=base["BodyText"], fontSize=9.5, leading=13),
    }
    styles["b1"] = ParagraphStyle("b1", parent=styles["body"], leftIndent=16)
    styles["b2"] = ParagraphStyle("b2", parent=styles["body"], leftIndent=32)
    SimpleDocTemplate(local_path, pagesize=LETTER, leftMargin=0.9 * inch, rightMargin=0.9 * inch,
                      topMargin=0.8 * inch, bottomMargin=0.8 * inch).build(_flowables(md, styles))


def _write_docx(local_path, md):
    from docx import Document
    d = Document()
    for raw in md.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("## "):
            d.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            d.add_heading(line[2:], level=1)
        elif line.strip().startswith("- "):
            ind = len(line) - len(line.lstrip())
            d.add_paragraph(line.strip()[2:], style="List Bullet 2" if ind >= 2 else "List Bullet")
        else:
            d.add_paragraph(line)
    d.save(local_path)


def main():
    # reportlab + python-docx are provided by the job environment
    # (resources/jobs.yml: wave1_synth_bronze + app_protocol_extract env deps).
    # We do NOT pip-install at runtime — a runtime install into a live serverless
    # kernel can hang/crash it (same hazard class as in-kernel w.postgres).
    vol = cfg().protocols_volume_path
    tmp = "/local_disk0/tmp/protocols" if os.path.isdir("/local_disk0") else "/tmp/protocols"
    os.makedirs(tmp, exist_ok=True)

    # poc_high as DOCX to exercise the Office path; the others as PDF.
    plan = {"poc_low": "pdf", "poc_med": "pdf", "poc_high": "docx"}
    written = []
    for sid, fmt in plan.items():
        local = f"{tmp}/{sid}.{fmt}"
        (_write_docx if fmt == "docx" else _write_pdf)(local, PROTOCOLS[sid])
        dst = f"{vol}/{sid}.{fmt}"
        shutil.copyfile(local, dst)
        written.append((dst, os.path.getsize(dst)))
    print(f"[seed] wrote {len(written)} protocol files to {vol}:")
    for p, n in written:
        print(f"  {p} ({n} bytes)")


if __name__ == "__main__":
    main()
